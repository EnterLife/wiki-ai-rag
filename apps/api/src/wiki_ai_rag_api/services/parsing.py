from __future__ import annotations

import csv
import json
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any


@dataclass(frozen=True)
class ParsedSegment:
    text: str
    metadata: dict[str, Any] = field(default_factory=dict)


def parse_file(path: Path) -> str:
    return "\n\n".join(segment.text for segment in parse_file_segments(path))


def parse_file_segments(path: Path) -> list[ParsedSegment]:
    suffix = path.suffix.lower()
    if suffix == ".md":
        return _parse_markdown(path)
    if suffix == ".txt":
        return [ParsedSegment(text=_read_text(path))]
    if suffix == ".csv":
        return _parse_csv(path)
    if suffix == ".json":
        return _parse_json(path)
    if suffix == ".html":
        return _parse_html(path)
    if suffix == ".pdf":
        return _parse_pdf(path)
    if suffix == ".docx":
        return _parse_docx(path)
    raise ValueError(f"Unsupported file extension: {suffix}")


def _read_text(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="cp1251")


def _parse_markdown(path: Path) -> list[ParsedSegment]:
    segments: list[ParsedSegment] = []
    section: str | None = None
    lines: list[str] = []
    for line in _read_text(path).splitlines():
        heading = re.match(r"^#{1,6}\s+(.+?)\s*$", line)
        if heading:
            if lines:
                segments.append(
                    ParsedSegment(
                        text="\n".join(lines).strip(),
                        metadata={"section": section} if section else {},
                    )
                )
            section = heading.group(1).strip()
            lines = [line]
        else:
            lines.append(line)
    if lines:
        segments.append(
            ParsedSegment(
                text="\n".join(lines).strip(),
                metadata={"section": section} if section else {},
            )
        )
    return [segment for segment in segments if segment.text]


def _parse_html(path: Path) -> list[ParsedSegment]:
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(_read_text(path), "html.parser")
    segments: list[ParsedSegment] = []
    section: str | None = None
    lines: list[str] = []
    for element in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "pre"]):
        text = element.get_text(" ", strip=True)
        if not text:
            continue
        if element.name.startswith("h"):
            if lines:
                segments.append(
                    ParsedSegment(
                        text="\n".join(lines),
                        metadata={"section": section} if section else {},
                    )
                )
            section = text
            lines = [text]
        else:
            lines.append(text)
    if lines:
        segments.append(
            ParsedSegment(
                text="\n".join(lines),
                metadata={"section": section} if section else {},
            )
        )
    return segments or [ParsedSegment(text=soup.get_text("\n"))]


def _parse_pdf(path: Path) -> list[ParsedSegment]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    return [
        ParsedSegment(text=text, metadata={"page": page_number})
        for page_number, page in enumerate(reader.pages, start=1)
        if (text := (page.extract_text() or "").strip())
    ]


def _parse_docx(path: Path) -> list[ParsedSegment]:
    from docx import Document

    document = Document(str(path))
    segments: list[ParsedSegment] = []
    section: str | None = None
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph.style and paragraph.style.name.lower().startswith("heading"):
            if lines:
                segments.append(
                    ParsedSegment(
                        text="\n".join(lines),
                        metadata={"section": section} if section else {},
                    )
                )
            section = text
            lines = [text]
        else:
            lines.append(text)
    table_lines = [
        " | ".join(cell.text.strip() for cell in row.cells)
        for table in document.tables
        for row in table.rows
    ]
    lines.extend(line for line in table_lines if line.strip(" |"))
    if lines:
        segments.append(
            ParsedSegment(
                text="\n".join(lines),
                metadata={"section": section} if section else {},
            )
        )
    return segments


def _parse_csv(path: Path) -> list[ParsedSegment]:
    with path.open("r", encoding="utf-8-sig", newline="") as csv_file:
        rows = list(csv.DictReader(csv_file))
    return [
        ParsedSegment(
            text="\n".join(f"{key}: {value}" for key, value in row.items()),
            metadata={"record_id": str(index)},
        )
        for index, row in enumerate(rows, start=1)
    ]


def _parse_json(path: Path) -> list[ParsedSegment]:
    payload = json.loads(_read_text(path))
    records = payload if isinstance(payload, list) else [payload]
    return [
        ParsedSegment(
            text=json.dumps(record, ensure_ascii=False, indent=2),
            metadata={"record_id": str(index)},
        )
        for index, record in enumerate(records, start=1)
    ]
